"""DOCX document loader."""

import asyncio
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

from app.ingestion.base_loader import BaseLoader, DocumentContent


class DocxLoader(BaseLoader):
    """Loader for DOCX documents."""
    
    supported_extensions = ["docx"]
    
    async def load(self, file_path: str | Path) -> DocumentContent:
        """Load and extract content from a DOCX file."""
        file_path = Path(file_path)
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self._load_sync, file_path)
    
    def _load_sync(self, file_path: Path) -> DocumentContent:
        """Synchronous DOCX loading."""
        try:
            doc = DocxDocument(file_path)
        except PackageNotFoundError as e:
            raise ValueError(f"Invalid DOCX file: {e}")
        
        # Extract paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append({
                    "text": text,
                    "style": para.style.name if para.style else None,
                })
        
        # Extract tables
        tables_text = []
        for table in doc.tables:
            table_content = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_content.append(" | ".join(row_text))
            tables_text.append("\n".join(table_content))
        
        # Combine all text
        full_text_parts = [p["text"] for p in paragraphs]
        full_text_parts.extend(tables_text)
        full_text = "\n\n".join(full_text_parts)
        
        # Extract metadata
        metadata = self._extract_metadata_sync(doc)
        metadata["file_name"] = file_path.name
        
        # Create pages (DOCX doesn't have natural pages, so we create sections)
        pages = self._create_sections(paragraphs)
        
        return DocumentContent(
            text=full_text,
            metadata=metadata,
            pages=pages,
        )
    
    async def extract_metadata(self, file_path: str | Path) -> dict[str, Any]:
        """Extract metadata from a DOCX file."""
        file_path = Path(file_path)
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None, self._extract_metadata_from_file, file_path
        )
    
    def _extract_metadata_from_file(self, file_path: Path) -> dict[str, Any]:
        """Extract metadata from DOCX file."""
        doc = DocxDocument(file_path)
        return self._extract_metadata_sync(doc)
    
    def _extract_metadata_sync(self, doc: DocxDocument) -> dict[str, Any]:
        """Extract metadata from DOCX document."""
        metadata: dict[str, Any] = {}
        
        core_props = doc.core_properties
        
        if core_props.title:
            metadata["title"] = core_props.title
        
        if core_props.author:
            metadata["author"] = core_props.author
        
        if core_props.subject:
            metadata["subject"] = core_props.subject
        
        if core_props.keywords:
            metadata["keywords"] = core_props.keywords
        
        if core_props.created:
            metadata["creation_date"] = str(core_props.created)
        
        if core_props.modified:
            metadata["modification_date"] = str(core_props.modified)
        
        if core_props.last_modified_by:
            metadata["last_modified_by"] = core_props.last_modified_by
        
        if core_props.category:
            metadata["category"] = core_props.category
        
        return metadata
    
    def _create_sections(
        self, paragraphs: list[dict[str, Any]]
    ) -> list[dict[str, Any]]:
        """Create sections from paragraphs based on headings."""
        pages = []
        current_section: list[str] = []
        current_title = ""
        page_num = 1
        
        heading_styles = ["Heading 1", "Heading 2", "Title"]
        
        for para in paragraphs:
            style = para.get("style", "")
            text = para["text"]
            
            if style in heading_styles and current_section:
                # Save current section
                pages.append({
                    "page_number": page_num,
                    "text": "\n".join(current_section),
                    "title": current_title,
                })
                page_num += 1
                current_section = [text]
                current_title = text
            else:
                if not current_title and style in heading_styles:
                    current_title = text
                current_section.append(text)
        
        # Save last section
        if current_section:
            pages.append({
                "page_number": page_num,
                "text": "\n".join(current_section),
                "title": current_title,
            })
        
        return pages if pages else [{"page_number": 1, "text": "", "title": ""}]
